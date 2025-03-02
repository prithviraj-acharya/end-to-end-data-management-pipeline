import os
import yaml
import boto3
import pandas as pd
from datetime import datetime
from io import StringIO
from docx import Document
from logs.logging_config import setup_logging  # Your logging configuration

# Setup logging
logger = setup_logging("ingestion.log")

# Get project root
current_dir = os.path.dirname(os.path.abspath(__file__))
project_root = os.path.abspath(os.path.join(current_dir, "../../"))

# Load AWS credentials from YAML
credentials_path = os.path.join(project_root, "config", "credentials.yaml")
with open(credentials_path, "r") as file:
    creds = yaml.safe_load(file)

bucket_name = "dmml-assignment1-bucket"


def generate_report(validation_details, source, file_name):
    """
    Generate a Word report with validation details.
    The report is saved under scripts/validation_reports.
    """
    report_dir = os.path.join(project_root, "scripts", "validation_reports")
    os.makedirs(report_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    report_file_name = f"validation_report_{source}_{file_name}_{timestamp}.docx"
    report_path = os.path.join(report_dir, report_file_name)

    doc = Document()
    doc.add_heading(f"Validation Report for {source} - {file_name}", 0)
    for key, value in validation_details.items():
        doc.add_paragraph(f"{key}: {value}")
    doc.save(report_path)
    logger.info(f"Data quality report generated: {report_path}")
    print(f"Validation report generated: {report_path}")


def validate_dataframe(data, source, file_name):
    """
    Validate a DataFrame and record detailed metrics.

    For 'local' data, expected validations are:
      - Numeric: 'tenure' (range 0–100), 'Churn' (range 0–1)
      - Categorical: 'gender' (["Male", "Female"]), 'PhoneService' (["Yes", "No"])

    For 'kaggle' data, expected validations are:
      - Numeric: 'tenure' (range 0–100), 'MonthlyCharges' (range 0–500), 'TotalCharges' (range 0–10000)
      - Categorical: 'gender' (["Male", "Female"]), 'InternetService' (["DSL", "Fiber optic", "No"])

    Also reports missing data and duplicates (based on 'customerID').
    Additional numeric columns (if not high-cardinality and not blacklisted) are validated for outliers.
    Additional categorical columns (except blacklisted) are summarized.
    """
    # Strip extra whitespace from column names (do not force lowercase)
    data.columns = data.columns.str.strip()

    validation_report = {}
    total_rows = len(data)

    # --- Missing Data Check ---
    missing_data = data.isnull().sum()
    missing_info = {col: f"{cnt} missing ({cnt / total_rows * 100:.2f}%)" for col, cnt in missing_data.items()}
    validation_report["missing_data"] = missing_info

    # Define expected validations based on source (use original case)
    if source == "local":
        expected_numeric = {"tenure": (0, 100), "Churn": (0, 1)}
        expected_categorical = {"gender": ["Male", "Female"], "PhoneService": ["Yes", "No"]}
    elif source == "kaggle":
        expected_numeric = {"tenure": (0, 100), "MonthlyCharges": (0, 500), "TotalCharges": (0, 10000)}
        expected_categorical = {"gender": ["Male", "Female"], "InternetService": ["DSL", "Fiber optic", "No"]}
    else:
        expected_numeric = {}
        expected_categorical = {}

    # --- Numeric Validations for expected columns ---
    for col, exp_range in expected_numeric.items():
        if col in data.columns:
            # Force conversion so expected columns are numeric
            data[col] = pd.to_numeric(data[col], errors="coerce")
            valid_data = data[col].dropna()
            if valid_data.nunique() <= 2:
                msg = "Binary column; outlier check not applicable."
            else:
                out_of_range = valid_data[(valid_data < exp_range[0]) | (valid_data > exp_range[1])]
                msg = f"{len(out_of_range)} out of {total_rows} rows ({len(out_of_range) / total_rows * 100:.2f}%) are outside the expected range {exp_range}."
            validation_report[col] = msg
        else:
            validation_report[col] = f"{col} column not found."

    # --- Numeric Validations for additional numeric columns ---
    numeric_blacklist = {"SeniorCitizen"}
    additional_numeric = set(data.select_dtypes(include=["number"]).columns) - set(
        expected_numeric.keys()) - numeric_blacklist
    for col in additional_numeric:
        if data[col].nunique() >= 0.9 * total_rows:
            validation_report[col] = f"High-cardinality numeric column; details omitted."
            continue
        col_numeric = pd.to_numeric(data[col], errors="coerce")
        valid_data = col_numeric.dropna()
        if valid_data.nunique() <= 2:
            msg = "Binary column; outlier check not applicable."
        else:
            Q1 = valid_data.quantile(0.25)
            Q3 = valid_data.quantile(0.75)
            IQR = Q3 - Q1
            lower_bound = Q1 - 1.5 * IQR
            upper_bound = Q3 + 1.5 * IQR
            outlier_count = valid_data[(valid_data < lower_bound) | (valid_data > upper_bound)].count()
            msg = f"{outlier_count} out of {total_rows} rows ({outlier_count / total_rows * 100:.2f}%) are outliers."
        validation_report[col] = msg

    # --- Categorical Validations for expected columns ---
    for col, expected_vals in expected_categorical.items():
        if col in data.columns:
            invalid = data[~data[col].isin(expected_vals)]
            count_invalid = len(invalid)
            msg = f"{count_invalid} invalid values out of {total_rows} rows ({count_invalid / total_rows * 100:.2f}%)."
            if count_invalid > 0:
                msg += f" Found: {invalid[col].unique().tolist()}."
            else:
                msg += " (All values valid)"
            validation_report[col] = msg
        else:
            validation_report[col] = f"{col} column not found."

    # --- Additional Categorical Columns ---
    blacklist = {"customerID"}
    additional_categorical = set(data.select_dtypes(include=["object"]).columns) - set(
        expected_categorical.keys()) - blacklist
    for col in additional_categorical:
        unique_vals = data[col].unique()
        if len(unique_vals) >= 0.9 * total_rows:
            validation_report[col] = f"High-cardinality column with {len(unique_vals)} unique values; details omitted."
        else:
            validation_report[col] = f"Unique values: {unique_vals.tolist()}"

    # --- Duplicates Check (on 'customerID') ---
    if "customerID" in data.columns:
        dup_count = data.duplicated(subset=["customerID"]).sum()
        validation_report[
            "duplicates"] = f"{dup_count} duplicate rows found out of {total_rows} rows ({dup_count / total_rows * 100:.2f}%)."
    else:
        validation_report["duplicates"] = "customerID column not found."

    # --- Outlier Detection for expected numeric columns ---
    outliers_info = {}
    for col in expected_numeric.keys():
        if col in data.columns:
            col_numeric = pd.to_numeric(data[col], errors="coerce")
            if col_numeric.nunique() <= 2:
                outliers_info[col] = "Binary column; outlier check not applicable."
                continue
            Q1 = col_numeric.quantile(0.25)
            Q3 = col_numeric.quantile(0.75)
            IQR = Q3 - Q1
            lower_bound = Q1 - 1.5 * IQR
            upper_bound = Q3 + 1.5 * IQR
            outlier_count = col_numeric[(col_numeric < lower_bound) | (col_numeric > upper_bound)].count()
            outliers_info[
                col] = f"{outlier_count} out of {total_rows} rows ({outlier_count / total_rows * 100:.2f}%) are outliers."
    for col in additional_numeric:
        # If details omitted due to high-cardinality, skip
        if col in validation_report and "High-cardinality" in validation_report[col]:
            continue
        col_numeric = pd.to_numeric(data[col], errors="coerce")
        if col_numeric.nunique() <= 2:
            outliers_info[col] = "Binary column; outlier check not applicable."
            continue
        Q1 = col_numeric.quantile(0.25)
        Q3 = col_numeric.quantile(0.75)
        IQR = Q3 - Q1
        lower_bound = Q1 - 1.5 * IQR
        upper_bound = Q3 + 1.5 * IQR
        outlier_count = col_numeric[(col_numeric < lower_bound) | (col_numeric > upper_bound)].count()
        outliers_info[
            col] = f"{outlier_count} out of {total_rows} rows ({outlier_count / total_rows * 100:.2f}%) are outliers."
    validation_report["outliers"] = outliers_info

    generate_report(validation_report, source, file_name)
    return validation_report


def validate_s3_files():
    """
    For each source ('kaggle' and 'local'), list objects in S3 under the given prefix,
    select the latest file (by LastModified), download its content, read it into a DataFrame,
    and validate it using source-specific expectations.
    """
    s3_client = boto3.client(
        "s3",
        aws_access_key_id=creds["aws"]["access_key"],
        aws_secret_access_key=creds["aws"]["secret_key"],
    )

    prefixes = {"kaggle": "kaggle/", "local": "local/"}

    for source, prefix in prefixes.items():
        logger.info(f"Listing objects for prefix: {prefix}")
        response = s3_client.list_objects_v2(Bucket=bucket_name, Prefix=prefix)
        objects = response.get("Contents", [])
        if objects:
            latest_obj = max(objects, key=lambda x: x["LastModified"])
            key = latest_obj["Key"]
            if key.endswith("/"):
                logger.info(f"Latest key for {source} is a folder. Skipping.")
                continue
            try:
                s3_object = s3_client.get_object(Bucket=bucket_name, Key=key)
                file_content = s3_object["Body"].read().decode("utf-8")
                data = pd.read_csv(StringIO(file_content))
                file_name = os.path.basename(key)
                logger.info(f"Validating latest {source} file: {file_name}")
                validate_dataframe(data, source, file_name)
                logger.info(f"Data quality report for {file_name} generated.")
            except Exception as e:
                logger.error(f"Error processing {key}: {str(e)}")
        else:
            logger.info(f"No objects found for prefix: {prefix}")


if __name__ == "__main__":
    validate_s3_files()
